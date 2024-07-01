from docx import Document
from docx.shared import Pt
from datetime import datetime
from tkinter import filedialog
import os
import comtypes.client

class CreateContract:

    def convertDocxPDF(self, indir, outdir):
        wdFormatPDF = 17

        in_file = os.path.abspath(indir)
        out_file = os.path.abspath(outdir)

        word = comtypes.client.CreateObject('Word.Application')
        word.visible = False
        doc = word.Documents.Open(in_file)
        doc.SaveAs(out_file, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()
        os.remove(in_file)

    # Capitalise first letter of each word but keep letters already capitalised
    def titleCustom(self, s):
        return ' '.join(word[0].upper() + word[1:] if word.isupper() else word.capitalize() for word in s.split())

    def replaceText(self, p, replaceList):
        for key, value in replaceList.items():
            key = "<{}>".format(key)
            if key in p.text:
                p.text = p.text.replace(key, value)

    def replaceParagraph(self, doc, replaceList):
        for p in doc.paragraphs:
            p.style = doc.styles['Normal']
            self.replaceText(p, replaceList)

    def replaceTable(self, doc, replaceList):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    p = cell.paragraphs[0]
                    p.style = doc.styles['Normal']
                    self.replaceText(p, replaceList)

    def createContract(self, template):
        style = template.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        replaceList = {
            "curr_date": self.currDate.strftime("%d %B %Y"),
            "TITLE": self.title.upper(),
            "FIRSTNAME": self.firstname.upper(),
            "SURNAME": self.surname.upper(),
            "STREET": self.street.upper(),
            "CITY": self.city.upper(),
            "STATE": self.state.upper(),
            "POSTCODE": self.postcode,
            "email": self.email,
            "firstname": self.titleCustom(self.firstname),
            "position": self.titleCustom(self.position),
            "employer": self.titleCustom(self.employer),
            "start_date": self.startDate.strftime("%A %d %B %Y"),
            "employment_type": self.employType.lower(),
            "tep": f"{self.tep:,.2f}",
            "salary": f"{self.baseSalary:,.2f}",
            "super": f"{self.superannuation:,.2f}",
            "manager": self.titleCustom(self.manager),
            "surname": self.titleCustom(self.surname)
        }

        self.replaceParagraph(template, replaceList)
        self.replaceTable(template, replaceList)

        # save the new document
        filename = "Employment Contract ({})".format(self.firstname.title() + " " + self.surname.title())
        directory = filedialog.askdirectory()
        
        tempfile = "temp_contract.docx"
        template.save(tempfile)

        if os.path.exists(f"{directory}/{filename}.pdf"):
            i = 1
            while os.path.exists(f"{directory}/{filename} ({i}).pdf"):
                i += 1
            self.convertDocxPDF(tempfile, f"{directory}/{filename} ({i}).pdf")
        else:
            self.convertDocxPDF(tempfile, f"{directory}/{filename}.pdf")

    def __init__(self, template, values):
        self.currDate = datetime.now().date()
        self.title = values["title"]
        self.firstname = values["firstname"]
        self.surname = values["surname"]
        self.street = values["street"]
        self.city = values["city"]
        self.state = values["state"]
        self.postcode = values["postcode"]
        self.email = values["email"]
        self.position = values["position"]
        self.employer = values["employer"]
        self.startDate = datetime.strptime(values["startDate"], '%d/%m/%Y')
        self.employType = values["employType"]
        self.tep = float(values["tep"])
        self.rate = float(values["rate"])
        self.baseSalary = self.tep / (1 + self.rate / 100)
        self.superannuation = self.tep - self.baseSalary
        self.manager = values["manager"]
        
        self.createContract(template)